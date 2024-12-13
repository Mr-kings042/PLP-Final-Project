import os

def generate_txt(title, content):
    file_path = f"{title}.txt"
    with open(file_path, "w") as file:
        file.write(f"Title: {title}\n\n{content}")
    return file_path

def generate_docx(title, content):
    from docx import Document

    file_path = f"{title}.docx"
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.save(file_path)
    return file_path
